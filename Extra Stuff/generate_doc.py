import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.color.rgb = RGBColor(0, 102, 204)

def add_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph(code)
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(50, 50, 50)
    # Add a slight background shading hack by using a border if needed, or just keep it simple.

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# --- Title Page ---
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph('Dungeon Escape\nComprehensive Technical Write-Up')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = 'Calibri'

doc.add_paragraph()
subtitle = doc.add_paragraph('Module: Programming in Java\nProject: 2D Tile-Based Action Game')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(16)
    run.font.name = 'Calibri'

doc.add_page_break()

# --- 1. Project Overview ---
add_heading(doc, '1. Project Overview', 1)
add_paragraph(doc, 'Dungeon Escape is a comprehensive 2D tile-based action-adventure game developed entirely using the Java programming language. The primary objective of the game requires the player to navigate through a series of procedurally generated dungeon environments, evade or engage hostile entities, collect resources (coins), and successfully locate and enter the exit portal to progress. The game is structured across 10 progressively challenging levels, culminating in a highly cinematic and intense boss battle followed by a "Void Chase" escape sequence in the final level.')
add_paragraph(doc, 'This project serves as a robust demonstration of advanced Java programming paradigms. It leverages both the JavaFX and Java Swing (AWT) graphical frameworks to deliver a seamless user experience. The application architecture relies heavily on object-oriented programming principles, multithreaded game loop synchronization, procedural generation algorithms, and complex artificial intelligence pathfinding. By integrating these disparate technologies, Dungeon Escape proves the versatility and power of Java for 2D game development without relying on external third-party game engines.')

try:
    doc.add_picture('screenshot_level1.png', width=Inches(5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, 'Figure 1: Dungeon Escape Level 1 Gameplay', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print(f"Could not add level1 image: {e}")

# --- 2. Architecture & Class Structure ---
add_heading(doc, '2. Architecture & Class Structure', 1)
add_paragraph(doc, 'The software architecture of Dungeon Escape is firmly rooted in a component-based design pattern. Rather than relying on deep, brittle inheritance hierarchies, the application utilizes composition, where distinct classes encapsulate specific responsibilities. The central coordinator is the GamePanel class, which acts as the core engine. It manages the game loop, delegates rendering tasks, and updates the state of all entity components.')
add_paragraph(doc, 'The graphical interface is bifurcated to utilize the strengths of two prominent Java UI toolkits:')
add_paragraph(doc, '• JavaFX: Deployed for the Main Menu system (GameMenu.java). It handles UI layouts, scalable vectors, CSS-styled buttons, and declarative animations. It runs on the JavaFX Application Thread.')
add_paragraph(doc, '• Java Swing / AWT: Deployed for the high-performance core game engine (GamePanel.java). It utilizes Graphics2D for direct pixel rendering, buffered image drawing, and runs its logic on a dedicated background thread.')

add_heading(doc, '2.1 Class Inventory and Responsibilities', 2)
add_paragraph(doc, 'The project comprises several focused classes, each handling a distinct subsystem:')
add_paragraph(doc, '• Main.java: The entry point of the application. It dictates whether the application launches directly into the game or initializes the JavaFX menu system.')
add_paragraph(doc, '• GameMenu.java: Implements the JavaFX application lifecycle. It manages scene transitions between the main title, mode selection, and the scoreboard, establishing dynamic property bindings to maintain aspect ratio across window resizes.')
add_paragraph(doc, '• GamePanel.java: The core game engine extending JPanel and implementing Runnable. It maintains the master state machine, executes the fixed-timestep loop, manages collision resolution, and orchestrates the rendering of the tile map and all entities.')
add_paragraph(doc, '• Player.java: Encapsulates all player-specific logic, including keyboard input interpretation, axis-aligned bounding box (AABB) collision, sprite animation handling, and win condition triggers.')
add_paragraph(doc, '• Enemy.java: The base enemy AI class. It handles autonomous movement, state tracking (wandering vs. agro), and utilizes the A* pathfinding algorithm to navigate the grid dynamically.')
add_paragraph(doc, '• ArcherBoss.java: An advanced adversarial entity introduced in later levels. It possesses complex state machines governing multi-phase attack patterns (cocking, burst firing, frenzy bullet-hell), procedural particle generation, and visual taunt systems.')
add_paragraph(doc, '• Map.java: Responsible for procedural dungeon generation. It manages the 2D integer array representing the grid, dynamically carves traversable paths, and ensures valid placement of hazards and the exit portal.')
add_paragraph(doc, '• ScoreManager.java: A utility class managing flat-file data persistence. It utilizes standard Java I/O streams to read and write high score entries to a local CSV file.')

try:
    doc.add_picture('screenshot_boss.png', width=Inches(5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, 'Figure 2: Level 5 Boss Encounter', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print(f"Could not add boss image: {e}")

# --- 3. Core Java Concepts Demonstrated ---
add_heading(doc, '3. Core Java Concepts Demonstrated', 1)

add_heading(doc, '3.1 Object-Oriented Programming (OOP) Principles', 2)
add_paragraph(doc, 'The entire architecture is a testament to the four foundational pillars of Object-Oriented Programming:')
add_paragraph(doc, 'Encapsulation: Internal states are strictly protected. For instance, the ArcherBoss class encapsulates its volatile combat parameters, such as frenzyTimer, burstPhase, and its collection of trailing ghost coordinates. The GamePanel interacts with the boss purely through public API methods like update() and draw(), ensuring the engine cannot inadvertently corrupt the boss’s internal state machine.')
add_paragraph(doc, 'Abstraction: Complex backend operations are hidden behind simplified interfaces. When GamePanel calls map.draw(Graphics2D g2), it remains entirely agnostic to the underlying mathematical logic required to calculate camera offsets, cull off-screen tiles, or map integer IDs to BufferedImage references. The engine simply requests the map to render itself.')
add_paragraph(doc, 'Inheritance and Polymorphism: The project utilizes inheritance to establish behavioral contracts. While the Player, Enemy, and ArcherBoss classes do not share a common custom superclass, they all conceptually adhere to an update/draw lifecycle. Furthermore, GamePanel inherits from javax.swing.JPanel, allowing it to hook into the Swing rendering pipeline by overriding the paintComponent() method.')
add_paragraph(doc, 'Composition: This is favored heavily over inheritance. The GamePanel class does not extend a hypothetical "GameWorld" class; instead, it is composed of a Map, a KeyHandler, a Player, and dynamic lists of Enemy and Projectile objects. This allows for immense flexibility, such as easily instantiating a completely different Map instance when moving to the Level 10 finale without tearing down the GamePanel.')

add_heading(doc, '3.2 Multithreading and the Game Loop', 2)
add_paragraph(doc, 'To prevent the game execution from freezing the Swing Event Dispatch Thread (EDT) and to ensure a consistent frame rate regardless of the host machine’s processing power, the game logic is decoupled into a dedicated background thread.')
add_paragraph(doc, 'The GamePanel implements the Runnable interface, defining a custom run() method. Within this method, a fixed-timestep game loop is executed using System.nanoTime() for high-resolution timing. The loop calculates a delta value based on a target of 60 Frames Per Second (FPS). Only when the delta accumulator exceeds 1.0 does the engine process physics (update()) and request a frame render (repaint()).')
add_code_block(doc, 'double drawInterval = 1000000000.0 / 60;\ndouble delta = 0;\nlong lastTime = System.nanoTime();\n\nwhile (gameThread != null) {\n    long currentTime = System.nanoTime();\n    delta += (currentTime - lastTime) / drawInterval;\n    lastTime = currentTime;\n    if (delta >= 1) {\n        update();\n        repaint();\n        delta--;\n    }\n}')
add_paragraph(doc, 'This delta-timing ensures that whether the host CPU is running at 1GHz or 5GHz, the player character will always move at exactly the intended speed per second.')

add_heading(doc, '3.3 Java Collections Framework', 2)
add_paragraph(doc, 'The java.util package is heavily leveraged to manage dynamic grouping of game entities. Unlike primitive arrays, which have fixed capacities, ArrayList is used extensively for entities that spawn and despawn dynamically. Active projectiles (Arrow.java), procedural particles, and trailing ghost coordinates are all maintained in ArrayList structures.')
add_paragraph(doc, 'Furthermore, the ScoreManager utilizes a List<ScoreEntry> to load and sort leaderboard data. Advanced Java features such as lambda expressions are utilized to implement custom comparators, sorting the scores primarily by points (descending) and secondarily by level reached:')
add_code_block(doc, 'scores.sort((a, b) -> {\n    if (b.score != a.score) return b.score - a.score;\n    return b.level - a.level;\n});')

add_heading(doc, '3.4 File I/O and Data Persistence', 2)
add_paragraph(doc, 'Persistent data storage is achieved using the java.io package. The ScoreManager class reads and writes to a local "highscores.txt" file using Comma-Separated Values (CSV) formatting. BufferedWriter and FileReader are wrapped in BufferedReader for optimized string tokenization.')
add_paragraph(doc, 'Modern Java features, specifically the try-with-resources statement, are employed to guarantee that file handles are securely closed and memory leaks are prevented, even in the event of an unexpected IOException during read/write operations.')

add_heading(doc, '3.5 Exception Handling', 2)
add_paragraph(doc, 'Robust error handling ensures the application remains stable. When loading external graphical assets (e.g., ImageIO.read()), the operations are wrapped in try-catch blocks. If a .png file is missing or corrupted, the catch block intercepts the exception and provides a graceful fallback—such as rendering a colored geometric primitive (e.g., filling a rectangle with Color.RED) instead of crashing the application with a NullPointerException.')

try:
    doc.add_picture('screenshot_gameover.png', width=Inches(5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, 'Figure 3: Game Over State and Score Prompt', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print(f"Could not add gameover image: {e}")

# --- 4. Advanced Algorithms and Logic ---
doc.add_page_break()
add_heading(doc, '4. Advanced Algorithms and Logic', 1)

add_heading(doc, '4.1 Procedural Map Generation', 2)
add_paragraph(doc, 'To ensure high replayability, the Map class dynamically generates a new dungeon layout for every level. The algorithm operates on a 2D integer array (int[][] map) and follows a structured procedure:')
add_paragraph(doc, '1. Initialization: The entire grid is populated with wall tiles (integer ID 1).')
add_paragraph(doc, '2. Room Carving: The algorithm selects random coordinates and carves out rectangular rooms by overriding the grid values with floor tiles (integer ID 0).')
add_paragraph(doc, '3. Path Enforcement: A path-carving algorithm guarantees that the map is solvable. It utilizes a greedy traversal, marching from the player’s spawn point to the designated exit point. At each step, it randomly selects to move horizontally or vertically toward the target, carving a 2x2 wide corridor to prevent claustrophobic bottlenecks.')
add_paragraph(doc, '4. Objective Placement: The exit portal is placed using Manhattan Distance calculations to ensure it spawns a mathematically enforced minimum distance away from the player.')

add_heading(doc, '4.2 A* (A-Star) Pathfinding Algorithm', 2)
add_paragraph(doc, 'Intelligent enemy navigation is driven by a custom implementation of the A* pathfinding algorithm. This ensures that enemies and bosses do not simply walk into walls when pursuing the player, but rather calculate the most efficient route around obstacles.')
add_paragraph(doc, 'The algorithm maps the grid to Node objects. Each Node evaluates three costs: gCost (exact distance from start), hCost (heuristic estimated distance to target via Manhattan calculation), and fCost (the sum of g and h). By maintaining an openList of nodes to evaluate and a closedList of evaluated nodes, the AI consistently expands the node with the lowest fCost. Once the target coordinate is reached, the optimal path is reconstructed by traversing the parent node references backward.')

add_heading(doc, '4.3 Collision Detection and Physics', 2)
add_paragraph(doc, 'Physics within the engine rely on Axis-Aligned Bounding Box (AABB) mathematics. The player’s rectangular bounds are continually verified against the integer grid array. The engine converts pixel coordinates into column/row indices using division by the tile size.')
add_paragraph(doc, 'To provide a smooth gameplay feel, horizontal and vertical collisions are resolved independently. If the player attempts to move diagonally into a wall, the engine will block the axis colliding with the wall while allowing movement along the free axis, simulating a "sliding" effect against the geometry.')

# --- 5. Graphics Rendering and Visual Effects ---
doc.add_page_break()
add_heading(doc, '5. Graphics Rendering and Visual Effects', 1)

add_heading(doc, '5.1 Custom Swing Rendering', 2)
add_paragraph(doc, 'The visual output of the game relies entirely on overriding the standard Swing paintComponent(Graphics g) pipeline. The base Graphics object is cast to a Graphics2D instance, unlocking advanced rendering capabilities such as sub-pixel translation, scaling matrices, and alpha blending.')

add_heading(doc, '5.2 Camera Translation and Adaptive Scaling', 2)
add_paragraph(doc, 'To support maps larger than the physical screen resolution, a dynamic camera system is implemented. Prior to drawing the map and entities, the rendering context is shifted using g2.translate(-cameraX, 0). This effectively moves the entire world in the opposite direction of the camera. After world rendering, the context is translated back so that HUD elements (like health and score) remain static on the screen.')
add_paragraph(doc, 'Additionally, adaptive scaling ensures the game looks correct on any monitor. The engine determines the smallest scale factor between the target logical resolution and the actual window size, applying g2.scale(scaleFactor, scaleFactor) to stretch the graphics uniformly while maintaining the correct aspect ratio.')

add_heading(doc, '5.3 Procedural Visual Effects', 2)
add_paragraph(doc, 'Dungeon Escape implements several high-fidelity visual effects procedurally, bypassing the need for heavy pre-rendered animations:')
add_paragraph(doc, '• Screen Shake: Triggered upon player damage or when proximity to the Void wall reaches critical levels. It is achieved by applying a randomized mathematical translation to the Graphics2D context immediately before the world rendering phase begins, causing the entire viewport to physically vibrate.')
add_paragraph(doc, '• Translucent Void Wall: The Level 10 chase sequence features a massive, consuming cosmic anomaly. This is rendered using a 60-layer iterative loop that calculates a smooth gradient transition from deep purple to pitch black. The alpha (transparency) values are dynamically adjusted, creating a glassmorphism effect that allows the underlying map geometry to remain visible beneath the cosmic fog. Furthermore, trigonometric functions (Math.sin()) are applied to the wall’s X-coordinate, creating a rhythmic, breathing pulse.')
add_paragraph(doc, '• Particle Systems: The ArcherBoss utilizes a procedural particle emitter during its attack phases. Points are spawned using random angles multiplied by Math.PI * 2, and their positions are calculated using sine and cosine functions to simulate inward charging energy and outward frenzy bursts.')

add_heading(doc, '5.4 JavaFX UI Binding and Animation', 2)
add_paragraph(doc, 'The GameMenu class exhibits advanced JavaFX capabilities. Scalability is achieved not by hardcoding sizes, but by establishing mathematical property bindings. The Scale transition pivot points are bound directly to the scene dimensions divided by the logical resolution, guaranteeing mathematically perfect aspect ratio preservation during live window resizing.')
add_paragraph(doc, 'Animations, such as the pulsating "PRESS ANY KEY" text, utilize JavaFX FadeTransition objects, operating asynchronously on the UI thread to provide a professional, polished introduction without manual frame-by-frame updates.')

# --- 6. State Management and Game Flow ---
doc.add_page_break()
add_heading(doc, '6. State Management and Game Flow', 1)
add_paragraph(doc, 'The overarching flow of the application is governed by a Finite State Machine (FSM) implemented via boolean flags within the GamePanel. Flags such as gameStarted, paused, gameOver, and gameWon dictate exactly which subset of the update() and draw() methods are executed during any given frame.')
add_paragraph(doc, 'The complexity of this state machine peaks during the Level 10 Finale. An integer variable, level10Phase, overrides standard procedural logic. When phase 1 is active, the engine locks the camera and enforces the destruction of Crystals and Bosses. When phase 2 is active, the engine swaps to a custom 300-tile chase map, initiates the Void Wall pursuit logic, and completely alters the win condition to rely on intersection with the final portal rather than a standard exit tile.')

try:
    doc.add_picture('screenshot_chase.png', width=Inches(5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, 'Figure 4: Level 10 Void Chase Sequence', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print(f"Could not add chase image: {e}")

# --- 7. Conclusion ---
add_heading(doc, '7. Conclusion', 1)
add_paragraph(doc, 'Dungeon Escape successfully synthesizes fundamental computer science concepts with advanced Java application design. By eschewing external game engines in favor of raw Java AWT/Swing and JavaFX implementations, the project demonstrates a profound understanding of the language’s core libraries, rendering pipelines, thread synchronization, and algorithmic problem-solving. From the recursive A* pathfinding calculations to the procedural alpha-blended visual effects, the codebase is a comprehensive portfolio piece showcasing end-to-end software engineering capability.')

doc.save('Dungeon_Escape_Technical_Writeup.docx')
